from docx.shared import Cm, Inches, Pt
from docx2pdf import convert
from matplotlib import style
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from docx.oxml.shape import CT_Picture
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml, register_element_cls
from docx.shared import RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.table import _Cell
from pyparsing import col
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENTATION
from docx import Document
from turtle import width
from pickle import NONE
from tkinter import CENTER
from turtle import color
import pandas as pd
import numpy as np
from datetime import datetime
import matplotlib.pyplot as plt
from math import pi
import matplotlib
import matplotlib.ticker as ticker

# 글꼴변경
matplotlib.rcParams['font.family'] = 'NanumGothic'

farmNameReport = input("농가명을 입력하세요. (입력 예시 : 김xx,이xx) : ")  # 농가명 입력
farmAdress = input("농장주소를 입력하세요. (주소 없을 시 Enter 클릭) : ")
farmName = list(farmNameReport.split(","))  # 농가명 입력
# farmName = 'test'


# 파일 업로드
def fileUpload(fileName):
    fileData = pd.read_excel(f'{fileName}.xlsx', dtype='object')
    return fileData


# 유전체 친자 업로드
parentageData = fileUpload('./data/parentage').replace('unknown', np.NaN)
parentageData = parentageData.astype({
    'IID': 'string',
    'KPN': 'string',
    'ACC': 'float'
})


# 유전체 대상우 리스트 업로드
cowList = fileUpload('./data/list').replace('unknown', np.NaN)
cowList = cowList.astype({
    'IID': 'string',
    'BIR': 'datetime64[ns]',
    'SEX': 'string',
    'FARM': 'string'
})


# 분석결과 업로드
referenceData = fileUpload('./data/reference').replace('unknown', np.NaN)
referenceData = referenceData.astype({
    'IID': 'string',
    'CWT_GEBV': 'float',
    'EMA_GEBV': 'float',
    'BFT_GEBV': 'float',
    'MAR_GEBV': 'float'
})

# 개월령 차트 구간 데이터 생성 함수


def countMonthDef(data, sex, min, max):
    countMonthData = data[(data['SEX'] == sex) & (
        (data['MONTHCOUNT'] >= min) & (data['MONTHCOUNT'] <= max))]
    countMonth = len(countMonthData)
    return countMonth

# 개월령 차트 생성 함수


def countMonthChart(dataframe, sex):
    plt.cla()
    plt.figure(figsize=(15, 11))
    plt.bar(dataframe['index'], dataframe[sex], color=['#EC7B2F'])

    for i, v in enumerate(dataframe['index']):

        plt.text(v, dataframe[sex][i], dataframe[sex][i],
                 fontsize=9,
                 color='black',
                 horizontalalignment='center',
                 verticalalignment='bottom')
        plt.savefig(f'./chart/monthchart_{sex}.jpg', bbox_inches='tight')
    plt.close()

# 표준화 함수


def sbv(x, mean, std):
    result = (x-mean)/std
    return result

# 선발지수 함수


def selectionIndex(x, weight):
    result = x * weight
    return result

# 순위백분율 함수


def pRank(x, max):
    result = x/max * 100
    return result

# 등급 함수


def grade(x):
    if x < 3:
        result = "SS"
    elif 3 <= x < 10:
        result = "S"
    elif 10 <= x < 20:
        result = "A"
    elif 20 <= x < 45:
        result = "B"
    elif 45 <= x < 70:
        result = "C"
    elif 70 <= x <= 100:
        result = "D"
    else:
        result = float("nan")
    return result

# 레이더 차트 함수


def raderchart(df, min, max, mingrid, maxgrid, ID):

    labels = df.columns[1:]
    num_labels = len(labels)

    angles = [x/float(num_labels)*(2*pi) for x in range(num_labels)]  # 각 등분점
    angles += angles[:1]  # 시작점으로 다시 돌아와야하므로 시작점 추가
    # my_palette = plt.cm.get_cmap("Set2", len(df.index))
    # mycolor = ['#FD4B04','#95A3A6'] #['전체','농가']
    mycolor = ['#00175A', '#ED7625']
    # myalpha=[0.8,0.15]
    mylinestyle = ['dashed', 'solid']
    fig = plt.figure(figsize=(5, 5))
    fig.set_facecolor('white')
    ax = fig.add_subplot(polar=True)
    for i, row in df.iterrows():
        color = mycolor[i]
        data = df.iloc[i].drop('Character').tolist()
        data += data[:1]

        ax.set_theta_offset(pi / 2)  # 시작점
        ax.set_theta_direction(-1)  # 그려지는 방향 시계방향

        phenoName = ['도체중', '등심', '등지방', '근내']
        plt.xticks(angles[:-1], phenoName, fontsize=13)  # x축 눈금 라벨
        ax.tick_params(axis='x', which='major', pad=15)  # x축과 눈금 사이에 여백을 준다.
        ax.set_rlabel_position(0)  # y축 각도 설정(degree 단위)
        ax.axes.yaxis.set_ticklabels([])
        plt.ylim(min-mingrid, max+maxgrid)

        ax.plot(angles, data, color=color, linewidth=3,
                linestyle=mylinestyle[i], label=row.Character)  # 레이더 차트 출력
        # ax.fill(angles, data, color=color, alpha=myalpha[i]) ## 도형 안쪽에 색을 채워준다.

    plt.legend(loc=(0.9, 0.9))
    plt.savefig(f'./chart/{ID}_rader.jpg', bbox_inches='tight')
    plt.close(fig)

# 순위 백분율 차트 함수


def percentRankChart(data, columnname, type, IID):
    plt.cla()
    plt.figure(figsize=(3, 7))
    plt.scatter(1, data[f'{columnname}_SGEBV_PRank'], s=60, color='#ED7625')
    ax = plt.subplot()
    plt.xlim([0, 2])
    plt.ylim([0, 100])
    ax.invert_yaxis()
    ax.xaxis.set_major_locator(ticker.MultipleLocator(1))
    # ax.yaxis.set_major_locator(ticker.MultipleLocator(5))
    xinout = [f'{type}']
    xindex = [1]
    plt.xticks(xindex, xinout)

    ax.yaxis.set_major_locator(ticker.MultipleLocator(10))
    yinout = ['0%', '10%', '20%', '30%', '40%',
              '50%', '60%', '70%', '80%', '90%', '100%',]
    yindex = [0, 10, 20, 30, 40, 50, 60, 70, 80, 90, 100]
    plt.yticks(yindex, yinout)

    plt.grid(True, alpha=0.5)
    plt.savefig(f'./chart/{IID}_{columnname}_pchart.jpg', bbox_inches='tight')
    plt.close()


# 대상농가 데이터 필터링
sortFarm = "(FARM in @farmName)"  # 쿼리문
filterCowList = cowList.query(sortFarm)  # 제작하고자 하는 농장에 해당하는 리스트만 필터링

# 농가현황(두수현황) table data 추출

# 두수현황 - 수
maleCowCount = filterCowList.loc[filterCowList['SEX'] == '수', 'SEX'].count()
# 두수현황 - 암
femaleCowCount = filterCowList.loc[filterCowList['SEX'] == '암', 'SEX'].count()
# 두수현황 - unknown
unknownCowCount = filterCowList['SEX'].isnull().sum(axis=0)
# 두수현황 - 전체
allCount = femaleCowCount + maleCowCount + unknownCowCount

sortFarm = "(FARM in @farmName)"  # 쿼리문
filterCowList = cowList.query(sortFarm)  # 제작하고자 하는 농장에 해당하는 리스트만 필터링

# 개월령 분포도 이미지 생성을 위한 데이터 추출

# 개월령 계산을 위한 datetime 분해
monthCountData = filterCowList.copy(deep=True)
monthCountData['NOW'] = datetime.now().date()
monthCountData['NOWYEAR'] = datetime.now().date().strftime('%y')
monthCountData['NOWMONTH'] = datetime.now().date().strftime('%m')
monthCountData['NOWDAY'] = datetime.now().date().strftime('%d')
monthCountData['BIRYEAR'] = monthCountData['BIR'].dt.strftime('%y')
monthCountData['BIRMONTH'] = monthCountData['BIR'].dt.strftime('%m')
monthCountData['BIRDAY'] = monthCountData['BIR'].dt.strftime('%d')

# 개월령 계산을 위한 type변경
monthCountData = monthCountData.astype({
    'NOWYEAR': 'float',
    'NOWMONTH': 'float',
    'NOWDAY': 'float',
    'BIRYEAR': 'float',
    'BIRMONTH': 'float',
    'BIRDAY': 'float'
})

# 개월령 계산 식
monthCountData['MONTHCOUNT'] = (((monthCountData['NOWYEAR'] - monthCountData['BIRYEAR'])
                                * 12) + monthCountData['NOWMONTH'] - monthCountData['BIRMONTH'])

monthCountData = monthCountData.astype({
    'MONTHCOUNT': 'float'
})

# 만 나이 cheack
monthCountData['Cheack'] = monthCountData['NOWDAY'] >= monthCountData['BIRDAY']

monthCountData.loc[(monthCountData['Cheack'] == True),
                   'MONTHCOUNT'] = monthCountData['MONTHCOUNT'] + 1

# 개월령 최대최소 구간 설정
minMax = [[0, 3], [4, 6], [7, 9], [10, 12], [13, 15], [16, 18],
          [19, 21], [22, 24], [25, 27], [28, 30], [31, 33], [34, 500]]

# 개월령 분포도 차트 dataframe 제작용 list 생성

fMonth = []
mMonth = []
for i in minMax:
    countM = countMonthDef(monthCountData, sex='암', min=i[0], max=i[1])
    fMonth.append(countM)

for i in minMax:
    countM = countMonthDef(monthCountData, sex='수', min=i[0], max=i[1])
    mMonth.append(countM)

# 개월령 분포도 차트 dataframe 생성

countMonthResult = []
countMonthHeader = ['0-3', '4-6', '7-9', '10-12', '13-15',
                    '16-18', '19-21', '22-24', '25-27', '28-30', '31-33', '34이상']
countMonthResult.append(fMonth)
countMonthResult.append(mMonth)

countMonthPd = pd.DataFrame(
    countMonthResult, columns=countMonthHeader, index=['암', '수'])
countMonthPd = countMonthPd.transpose().reset_index()

# 개월령 분포도 차트 생성
countMonthChart(countMonthPd, sex='암')
countMonthChart(countMonthPd, sex='수')

# 레퍼런스 계산시트 제작
# 표준화 작업
referenceStd = referenceData.std(numeric_only=True)  # 표준편차구하기
referenceMean = referenceData.mean(numeric_only=True)  # 평균구하기

referenceStdData = referenceData.copy(deep=True)  # referenceData copy

referenceStdData['CWT_SGEBV'] = referenceStdData['CWT_GEBV'].map(
    lambda x: sbv(x, referenceMean['CWT_GEBV'], referenceStd['CWT_GEBV']))  # 도체중 표준화
referenceStdData['EMA_SGEBV'] = referenceStdData['EMA_GEBV'].map(lambda x: sbv(
    x, referenceMean['EMA_GEBV'], referenceStd['EMA_GEBV']))  # 등심단면적 표준화
referenceStdData['BFT_SGEBV'] = referenceStdData['BFT_GEBV'].map(lambda x: sbv(
    x, referenceMean['BFT_GEBV'], referenceStd['BFT_GEBV'])) * -1  # 등지방 표준화, 인버스처리
referenceStdData['MAR_SGEBV'] = referenceStdData['MAR_GEBV'].map(lambda x: sbv(
    x, referenceMean['MAR_GEBV'], referenceStd['MAR_GEBV']))  # 근내지방도 표준화

referenceSbvMean = referenceStdData.mean(numeric_only=True)  # 표준화 평균(차트용)

# 표준화 육종가 순위 구하기
referenceStdData['CWT_SGEBV_Rank'] = referenceStdData['CWT_SGEBV'].rank(
    method='dense', ascending=False)
referenceStdData['EMA_SGEBV_Rank'] = referenceStdData['EMA_SGEBV'].rank(
    method='dense', ascending=False)
referenceStdData['BFT_SGEBV_Rank'] = referenceStdData['BFT_SGEBV'].rank(
    method='dense', ascending=False)
referenceStdData['MAR_SGEBV_Rank'] = referenceStdData['MAR_SGEBV'].rank(
    method='dense', ascending=False)

# 육종가 순위 백분율 분모 산출
CWT_SGEBV_Rank_Max = referenceStdData['CWT_SGEBV_Rank'].max()
EMA_SGEBV_Rank_Max = referenceStdData['EMA_SGEBV_Rank'].max()
BFT_SGEBV_Rank_Max = referenceStdData['BFT_SGEBV_Rank'].max()
MAR_SGEBV_Rank_Max = referenceStdData['MAR_SGEBV_Rank'].max()

# 육종가 순위 백분율 데이터 생성
referenceStdData['CWT_SGEBV_PRank'] = referenceStdData['CWT_SGEBV_Rank'].map(
    lambda x: pRank(x, CWT_SGEBV_Rank_Max))
referenceStdData['EMA_SGEBV_PRank'] = referenceStdData['EMA_SGEBV_Rank'].map(
    lambda x: pRank(x, EMA_SGEBV_Rank_Max))
referenceStdData['BFT_SGEBV_PRank'] = referenceStdData['BFT_SGEBV_Rank'].map(
    lambda x: pRank(x, BFT_SGEBV_Rank_Max))
referenceStdData['MAR_SGEBV_PRank'] = referenceStdData['MAR_SGEBV_Rank'].map(
    lambda x: pRank(x, MAR_SGEBV_Rank_Max))

# 육종가 등급 데이터 생성
referenceStdData['CWT_SGEBV_grade'] = referenceStdData['CWT_SGEBV_PRank'].map(
    lambda x: grade(x))
referenceStdData['EMA_SGEBV_grade'] = referenceStdData['EMA_SGEBV_PRank'].map(
    lambda x: grade(x))
referenceStdData['BFT_SGEBV_grade'] = referenceStdData['BFT_SGEBV_PRank'].map(
    lambda x: grade(x))
referenceStdData['MAR_SGEBV_grade'] = referenceStdData['MAR_SGEBV_PRank'].map(
    lambda x: grade(x))

# 육종가
referenceStdData['Si'] = referenceStdData['CWT_SGEBV'].map(lambda x: selectionIndex(x, 7)) + referenceStdData['EMA_SGEBV'].map(lambda x: selectionIndex(
    x, 3)) + referenceStdData['BFT_SGEBV'].map(lambda x: selectionIndex(x, -3)) + referenceStdData['MAR_SGEBV'].map(lambda x: selectionIndex(x, 4))
referenceStdData['Si_Rank'] = referenceStdData['Si'].rank(
    method='dense', ascending=False)

Si_Rank_Max = referenceStdData['Si_Rank'].max()

referenceStdData['Si_PRank'] = referenceStdData['Si_Rank'].map(
    lambda x: pRank(x, Si_Rank_Max))


# RawDataFrame 생성
kpnMergeCowList = pd.merge(
    filterCowList, parentageData, how='left', on='IID')  # 필터링된 리스트랑 KPN merge
genomeReportRawData = pd.merge(
    kpnMergeCowList, referenceStdData, how='left', on='IID')  # 보고서 제작용 raw Data

# 개체별 page 데이터,차트 만들기
# 개체별 page dataframe생성
genomeReportData = genomeReportRawData.copy(deep=True)
idvCowData = genomeReportData.iloc[:, [
    0, 1, 2, 4, 5, 6, 7, 8, 9, 22, 23, 24, 25, 28]]

idvCowData = idvCowData.sort_values(by=['Si_PRank'])
idvCowData['ACC'] = idvCowData['ACC'].round(3) * 100
idvCowData['ACC'] = idvCowData['ACC'].astype(str) + '%'
idvCowData['BIR'] = idvCowData['BIR'].dt.date

idvCowData = idvCowData.fillna('unknown').round(2)
idvCowData['Si_PRank'] = idvCowData['Si_PRank'].astype(str) + '%'

# 개체별 page 차트 제작용 dataframe 생성

# 레이더 차트용 dataframe
genomeChartData = genomeReportRawData.copy(deep=True)
genomeChartData = genomeChartData.iloc[:, [0, 10, 11, 12, 13]]

# 순위백분율 차트용 dataframe
percentRankData = genomeReportRawData.copy(deep=True).round(2)
percentRankData = percentRankData[[
    'IID', 'CWT_SGEBV_PRank', 'EMA_SGEBV_PRank', 'BFT_SGEBV_PRank', 'MAR_SGEBV_PRank']]

# 데이터 프레임 길이
dataLength = len(genomeChartData.index)

# 레이더 차트 생성
print("유전능력 차트 생성 중 .....")
for i in range(dataLength):

    idvNm = genomeChartData.iloc[i]['IID']  # 해당 개체

    data = genomeChartData.iloc[[i], 1:].transpose()

    breedChartSeries = pd.concat([referenceSbvMean, data], axis=1)
    breedSeedDf = pd.DataFrame(
        index=['CWT_SGEBV', 'EMA_SGEBV', 'BFT_SGEBV', 'MAR_SGEBV'])
    breedChartDf = pd.concat([breedSeedDf, breedChartSeries], axis=1)
    breedChartDf.columns = ['전체평균', '개체평균']
    breedChartTable = breedChartDf.loc[[
        'CWT_SGEBV', 'EMA_SGEBV', 'BFT_SGEBV', 'MAR_SGEBV'], ['전체평균', '개체평균']]  # 차트용(SBV)

    breedChartTable = breedChartTable.transpose()
    breedChartTable = breedChartTable.reset_index()
    breedChartTable.columns = [
        'Character', 'CWT_SGEBV', 'EMA_SGEBV', 'BFT_SGEBV', 'MAR_SGEBV']

    breedRaderMax = breedChartTable.max(numeric_only=True, axis=1)
    breedRaderMin = breedChartTable.min(numeric_only=True, axis=1)
    breedRaderSum = pd.concat([breedRaderMax, breedRaderMin], axis=0)
    breedRaderSumValue = breedRaderSum.values
    breedRaderMaxValue = np.max(breedRaderSumValue)
    breedRaderMinValue = np.min(breedRaderSumValue)

    breedRaderChart = raderchart(breedChartTable, breedRaderMinValue,
                                 breedRaderMaxValue, mingrid=0.5, maxgrid=0.5, ID=idvNm)  # 레이더 차트 생성

print("순위백분율 차트 생성 중 .....")
# 순위백분율 차트 생성
for i in range(dataLength):
    rankIdvNm = percentRankData.iloc[i]['IID']  # 해당 개체

    rankData = percentRankData.iloc[[i], 1:]
    rankData.index = ['개체평균']
    rankData = rankData.reset_index()

    percentRankChartCwt = percentRankChart(
        rankData, 'CWT', '도체중', IID=rankIdvNm)
    percentRankChartEma = percentRankChart(
        rankData, 'EMA', '등심', IID=rankIdvNm)
    percentRankChartBft = percentRankChart(
        rankData, 'BFT', '등지방', IID=rankIdvNm)
    percentRankChartMar = percentRankChart(
        rankData, 'MAR', '근내', IID=rankIdvNm)
# 개체별 page 데이터,차트 만들기 완료

# 선발지수 적용 리스트
selectIndexData = idvCowData.copy(deep=True)
selectIndexData = selectIndexData.iloc[:, [0, 3, 1, 2, 5, 6, 7, 8, 13]]
selectIndexData.columns = ['개체번호', 'KPN', '생년월일',
                           '성별', '도체중', '등심단면적', '등지방두께', '근내지방', '순위']

# 선발지수 dataframe 생성

################# 보고서 제작#################보고서 제작#################보고서 제작#################보고서 제작#################

# 보고서 제작 class,def 모음

# Cm,Inches,Pt 단위를 사용하기 위한 모듈

# 문자 스타일 변경

# table border style

# para 정렬

# table 정렬


# font color

# -*- coding: utf-8 -*-

# filename: add_float_picture.py

'''
Implement floating image based on python-docx.
- Text wrapping style: BEHIND TEXT <wp:anchor behindDoc="1">
- Picture position: top-left corner of PAGE `<wp:positionH relativeFrom="page">`.
Create a docx sample (Layout | Positions | More Layout Options) and explore the
source xml (Open as a zip | word | document.xml) to implement other text wrapping
styles and position modes per `CT_Anchor._anchor_xml()`.
'''


# 사진 사이즈 및 위치 변경
# refer to docx.oxml.shape.CT_Inline

class CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapNone/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % (nsdecls('wp', 'a', 'pic',
                                      'r'), int(pos_x), int(pos_y))
        )


# refer to docx.parts.story.BaseStoryPart.new_pic_inline
def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    """Return a newly-created `w:anchor` element.
    The element contains the image specified by *image_descriptor* and is scaled
    based on the values of *width* and *height*.
    """
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)


# refer to docx.text.run.add_picture
def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page.
    """
    run = p.add_run()
    anchor = new_pic_anchor(run.part, image_path_or_stream,
                            width, height, pos_x, pos_y)
    run._r.add_drawing(anchor)


# refer to docx.oxml.__init__.py
register_element_cls('wp:anchor', CT_Anchor)

# 문단 텍스트 입력


def paragraphText(paragraph, text, fontsize, color, alignment, style):
    text = paragraph.add_run(text)
    text.font.size = Pt(fontsize)
    font = text.font
    font.color.rgb = RGBColor.from_string(color)
    if alignment == 'CENTER':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        pass
    if style == 'bold':
        text.bold = True
    else:
        pass


def lineSpace(doc, inches, space_before, space_after):
    lineSpace = doc.add_paragraph()
    lineSpace.paragraph_format.line_spacing = Inches(inches)
    lineSpace.paragraph_format.space_before = Pt(space_before)
    lineSpace.paragraph_format.space_after = Pt(space_after)

# table 제작


def makeTable(paragraph, row, col, alignment, width, height):

    table = paragraph.add_table(rows=row, cols=col)
    if width == None:
        pass
    else:
        set_col_widths(table, width)

    if height == None:
        pass
    else:
        set_col_height(table, height)

    if alignment == 'CENTER':
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    else:
        pass
    return table

# table 가로 크기 변경


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

# table 세로 크기 변경


def set_col_height(table, heights):
    for idx, row in enumerate(table.rows):
        row.height = heights[idx]

# table 전체 스타일 변경


def titleBorder(
    table,
    top_val,
    top_color,
    top_sz,
    bottom_val,
    bottom_color,
    bottom_sz,
    left_val,
    left_color,
    left_sz,
    right_val,
    right_color,
    right_sz
):

    tbl = table._tbl  # get xml element in table
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr  # get tcPr element, in which we can define style of borders
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), top_val)
        top.set(qn('w:color'), top_color)
        top.set(qn('w:sz'), top_sz)

        left = OxmlElement('w:left')
        left.set(qn('w:val'), left_val)
        left.set(qn('w:color'), left_color)
        left.set(qn('w:sz'), left_sz)

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), bottom_val)
        bottom.set(qn('w:color'), bottom_color)
        bottom.set(qn('w:sz'), bottom_sz)

        right = OxmlElement('w:right')
        right.set(qn('w:val'), right_val)
        right.set(qn('w:color'), right_color)
        right.set(qn('w:sz'), right_sz)

        tcBorders.append(top)
        tcBorders.append(left)
        tcBorders.append(bottom)
        tcBorders.append(right)
        tcPr.append(tcBorders)

# table border 부분 변경
# def set_cell_border(cell: _Cell, **kwargs):


def set_cell_border(table, row, col, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    cell = table.rows[row].cells[col]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

# 셀 개별 배경색 변환


def cellBackColor(table, row, cell, color):
    # GET CELLS XML ELEMENT
    cell_xml_element = table.rows[row].cells[cell]._tc
    # RETRIEVE THE TABLE CELL PROPERTIES
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    # CREATE SHADING OBJECT
    shade_obj = OxmlElement('w:shd')
    # SET THE SHADING OBJECT
    shade_obj.set(qn('w:fill'), color)
    # APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
    table_cell_properties.append(shade_obj)

# 셀 텍스트 삽입


def insertTextCell(table, row, col, text, color, fontSize, fontStyle, vertical_alignment, para_alignment, space_before, space_after):
    cell = table.rows[row].cells[col]
    paragraph = cell.paragraphs[0]
    inputText = paragraph.add_run(text)
    font = inputText.font
    font.color.rgb = RGBColor.from_string(color)

    if fontStyle == 'bold':
        inputText.bold = True
    else:
        pass

    if vertical_alignment == 'CENTER':
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    elif vertical_alignment == 'TOP':
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    elif vertical_alignment == 'BOTTOM':
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    elif vertical_alignment == 'BOTH':
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTH
    else:
        pass

    if para_alignment == 'CENTER':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif para_alignment == 'RIGHT':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif para_alignment == 'LEFT':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        pass

    paragraph.paragraph_format.space_before = space_before
    paragraph.paragraph_format.space_after = space_after

    font.size = Pt(fontSize)

# 셀 병합


def cellMerge(table, stdCellList, cellList):
    stdCell = table.cell(stdCellList[0], stdCellList[1])
    for i in cellList:
        mergeCell = table.cell(i[0], i[1])
        stdCell.merge(mergeCell)

# header 박스 만들기


def makeHeaderBox(No, text):
    headerBoxWidth = [Cm(1.5), Cm(1), Cm(24.5)]
    headerBoxHeight = [Cm(1)]
    headerBox = makeTable(doc, row=1, col=3, alignment='CENTER',
                          width=headerBoxWidth, height=headerBoxHeight)
    set_cell_border(
        headerBox,
        row=0,
        col=2,
        top={"val": "nil"},
        bottom={"val": "single", "sz": "15", "color": "#F58D22"},
        start={"val": "nil"},
        end={"val": "nil"}
    )

    insertTextCell(
        headerBox,
        row=0,
        col=0,
        text=No,
        color='FFFFFF',
        fontSize=20,
        fontStyle='bold',
        vertical_alignment='CENTER',
        para_alignment='CENTER',
        space_after=Pt(0),
        space_before=Pt(0)
    )
    insertTextCell(
        headerBox,
        row=0,
        col=2,
        text=text,
        color='000000',
        fontSize=20,
        fontStyle='bold',
        vertical_alignment='CENTER',
        para_alignment='LEFT',
        space_after=Pt(0),
        space_before=Pt(0)
    )

    cellBackColor(headerBox, row=0, cell=0, color="F58D22")

# header 반복 함수


def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row


# 문서 생성
doc = Document()

# 문서 전체 폰트 변경
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 전체 페이지 가로세로 설정
current_section = doc.sections[-1]

current_section.orientation = WD_ORIENTATION.LANDSCAPE
current_section.page_width = Cm(29.7)
current_section.page_height = Cm(21.0)

current_section.top_margin = Cm(1)
current_section.bottom_margin = Cm(1)
current_section.left_margin = Cm(1)
current_section.right_margin = Cm(1)

# page1 표지 만들기

sign = doc.add_paragraph()

# 사진의 크기를 Cm 단위로 설정하여 삽입
add_float_picture(sign, './signPicture/page1.png',
                  width=Cm(27.7), height=Cm(19), pos_x=Cm(1), pos_y=Cm(1))

lineSpace(doc, inches=0.6, space_before=0, space_after=0)
lineSpace(doc, inches=0.6, space_before=0, space_after=0)
lineSpace(doc, inches=0.6, space_before=0, space_after=0)

paraTitle1 = doc.add_paragraph()
title1 = paragraphText(paraTitle1, text='한우', fontsize=32,
                       color='000000', alignment='CENTER', style='bold')

paraTitle2 = doc.add_paragraph()
title2 = paragraphText(paraTitle2, text='유전체 보고서', fontsize=32,
                       color='000000', alignment='CENTER', style='bold')

paraTitle3 = doc.add_paragraph()
title3 = paragraphText(paraTitle3, text=f'{farmNameReport} 농장',
                       fontsize=18, color='000000', alignment='CENTER', style='bold')

# page1 표지 만들기 완료

# page2 목차 만들기
doc.add_page_break()

# 목차 만들기
lineSpace(doc, inches=0.3, space_before=0, space_after=0)

indexBoxWidth = [Cm(26)]
indexBoxHeight = [Cm(17)]
indexBox = makeTable(doc, row=1, col=1, alignment='CENTER',
                     width=indexBoxWidth, height=indexBoxHeight)

# 목차 테두리 완성
titleBorder(
    indexBox,
    top_val='single',
    top_color='#F58D22',
    top_sz='25',
    bottom_val='single',
    bottom_color='#F58D22',
    bottom_sz='25',
    left_val='single',
    left_color='#F58D22',
    left_sz='25',
    right_val='single',
    right_color='#F58D22',
    right_sz='25'
)

# title 박스 생성
indexBoxCell0_0 = indexBox.rows[0].cells[0]

titleBoxWidth = [Cm(17)]
titleBoxHeights = [Cm(0.5), Cm(0.5)]
titleBox = makeTable(indexBoxCell0_0, row=2, col=1,
                     alignment='CENTER', width=titleBoxWidth, height=titleBoxHeights)
titleBoxCell0_0 = titleBox.rows[0].cells[0]

# title 박스 테두리 변경
set_cell_border(
    titleBox,
    row=0,
    col=0,
    top={"val": "nil"},
    bottom={"val": "single", "sz": "20", "color": "#F58D22"},
    start={"val": "nil"},
    end={"val": "nil"},
)

# title 박스 텍스트 입력
insertTextCell(
    titleBox,
    row=0,
    col=0,
    text='유전체 보고서',
    color='000000',
    fontSize=13,
    fontStyle='bold',
    vertical_alignment='CENTER',
    para_alignment='RIGHT',
    space_after=Pt(0),
    space_before=Pt(0)
)

insertTextCell(
    titleBox,
    row=1,
    col=0,
    text='C',
    color='843C0C',
    fontSize=20,
    fontStyle='bold',
    vertical_alignment='CENTER',
    para_alignment='LEFT',
    space_after=Pt(0),
    space_before=Pt(0)
)

insertTextCell(
    titleBox,
    row=1,
    col=0,
    text='ONTENTS',
    color='F58D22',
    fontSize=20,
    fontStyle='bold',
    vertical_alignment='CENTER',
    para_alignment='LEFT',
    space_after=Pt(0),
    space_before=Pt(0)
)

# 목차 리스트 table 생성
titleListWidth = [Cm(2), Cm(8)]
titleListHeights = [Cm(1), Cm(0.5), Cm(1), Cm(0.5),
                    Cm(1), Cm(0.5), Cm(1), Cm(0.5)]
titleList = makeTable(indexBoxCell0_0, row=8, col=2, alignment='CENTER',
                      width=titleListWidth, height=titleListHeights)

# 목차 리스트 table cell 병합
stdCellList1 = [1, 0]
cellList1 = [[1, 1]]
cellMerge(titleList, stdCellList1, cellList1)

stdCellList2 = [3, 0]
cellList2 = [[3, 1]]
cellMerge(titleList, stdCellList2, cellList2)

stdCellList3 = [5, 0]
cellList3 = [[5, 1]]
cellMerge(titleList, stdCellList3, cellList3)

stdCellList4 = [7, 0]
cellList4 = [[7, 1]]
cellMerge(titleList, stdCellList4, cellList4)

# 목차 리스트 table 텍스트 입력
insertTextCell(
    titleList,
    row=0,
    col=0,
    text='01',
    color='F58D22',
    fontSize=25,
    fontStyle='bold',
    vertical_alignment='CENTER',
    para_alignment='CENTER',
    space_after=Pt(0),
    space_before=Pt(0)
)

insertTextCell(
    titleList,
    row=0,
    col=1,
    text=' 농가 정보',
    color='000000',
    fontSize=15,
    fontStyle='bold',
    vertical_alignment='CENTER',
    para_alignment=None,
    space_after=Pt(0),
    space_before=Pt(0)
)

insertTextCell(
    titleList,
    row=2,
    col=0,
    text='02',
    color='F58D22',
    fontSize=25,
    fontStyle='bold',
    vertical_alignment='CENTER',
    para_alignment='CENTER',
    space_after=Pt(0),
    space_before=Pt(0)
)

insertTextCell(
    titleList,
    row=2,
    col=1,
    text=' 개월령 별 개체 분포도',
    color='000000',
    fontSize=15,
    fontStyle='bold',
    vertical_alignment='CENTER',
    para_alignment=None,
    space_after=Pt(0),
    space_before=Pt(0)
)

insertTextCell(
    titleList,
    row=4,
    col=0,
    text='03',
    color='F58D22',
    fontSize=25,
    fontStyle='bold',
    vertical_alignment='CENTER',
    para_alignment='CENTER',
    space_after=Pt(0),
    space_before=Pt(0)
)

insertTextCell(
    titleList,
    row=4,
    col=1,
    text=' 개체별 유전능력',
    color='000000',
    fontSize=15,
    fontStyle='bold',
    vertical_alignment='CENTER',
    para_alignment=None,
    space_after=Pt(0),
    space_before=Pt(0)
)

insertTextCell(
    titleList,
    row=6,
    col=0,
    text='04',
    color='F58D22',
    fontSize=25,
    fontStyle='bold',
    vertical_alignment='CENTER',
    para_alignment='CENTER',
    space_after=Pt(0),
    space_before=Pt(0)
)

insertTextCell(
    titleList,
    row=6,
    col=1,
    text=' 농가 유전능력 순위[선발지수]',
    color='000000',
    fontSize=15,
    fontStyle='bold',
    vertical_alignment='CENTER',
    para_alignment=None,
    space_after=Pt(0),
    space_before=Pt(0)
)

# 목차 리스트 table 테두리 변경
set_cell_border(
    titleList,
    row=0,
    col=0,
    top={"val": "nil"},
    bottom={"val": "nil"},
    start={"val": "nil"},
    end={"val": "single", "sz": "15"},
)

set_cell_border(
    titleList,
    row=2,
    col=0,
    top={"val": "nil"},
    bottom={"val": "nil"},
    start={"val": "nil"},
    end={"val": "single", "sz": "15"},
)

set_cell_border(
    titleList,
    row=4,
    col=0,
    top={"val": "nil"},
    bottom={"val": "nil"},
    start={"val": "nil"},
    end={"val": "single", "sz": "15"},
)

set_cell_border(
    titleList,
    row=6,
    col=0,
    top={"val": "nil"},
    bottom={"val": "nil"},
    start={"val": "nil"},
    end={"val": "single", "sz": "15"},
)

doc.add_page_break()

# page2 목차 만들기 완료

# page3 농가 현황 만들기

# 01.농가정보 headerBox 생성
makeHeaderBox(No="01", text="농가 정보")

lineSpace(doc, inches=0.1, space_before=0, space_after=0)

# 농가현황 컨테이너 생성
farmInfoContainerWidth = [Cm(13.5), Cm(13.5)]
farmInfoContainerHeight = [Cm(6)]
farmInfoContainer = makeTable(doc, row=1, col=2, alignment='CENTER',
                              width=farmInfoContainerWidth, height=farmInfoContainerHeight)

# 농가현황(농장정보) table 단락 공백 조절
fInfocontainerCell0_0 = farmInfoContainer.rows[0].cells[0]

farmInfoContainerParagraph0_0 = farmInfoContainer.rows[0].cells[0].paragraphs[0]
farmInfoContainerParagraph0_0.paragraph_format.line_spacing = Inches(0.01)
farmInfoContainerParagraph0_0.paragraph_format.space_before = Pt(0)
farmInfoContainerParagraph0_0.paragraph_format.space_after = Pt(0)

# 농가현황(농장정보) table 생성
farmInfoBoxWidth = [Cm(2), Cm(11.5)]
farmInfoBoxHeight = [Cm(2.7), Cm(2.7)]
farmInfoBox = makeTable(fInfocontainerCell0_0, row=2, col=2,
                        alignment='CENTER', width=farmInfoBoxWidth, height=farmInfoBoxHeight)

# 농가현황(농장정보) table 테두리 변경
titleBorder(
    farmInfoBox,
    top_val='single',
    top_color='#F58D22',
    top_sz='5',
    bottom_val='single',
    bottom_color='#F58D22',
    bottom_sz='5',
    left_val='single',
    left_color='#F58D22',
    left_sz='5',
    right_val='single',
    right_color='#F58D22',
    right_sz='5'
)

# 농가현황(농장정보) table 배경색 변경
cellBackColor(farmInfoBox, row=0, cell=0, color='FDECDA')
cellBackColor(farmInfoBox, row=1, cell=0, color='FDECDA')

# 농가현황(농장정보) table index 입력
insertTextCell(farmInfoBox, 0, 0, '이 름', '000000', fontSize=12, fontStyle='bold',
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
insertTextCell(farmInfoBox, 1, 0, '위 치', '000000', fontSize=12, fontStyle='bold',
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

# 농가현황(농장정보) table data 입력
insertTextCell(farmInfoBox, 0, 1, farmNameReport, '000000', fontSize=12, fontStyle=None,
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
insertTextCell(farmInfoBox, 1, 1, farmAdress, '000000', fontSize=12, fontStyle=None,
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))


# 농가현황(두수현황) table 단락 공백 조절
fInfocontainerCell0_1 = farmInfoContainer.rows[0].cells[1]

farmInfoContainerParagraph0_1 = farmInfoContainer.rows[0].cells[1].paragraphs[0]
farmInfoContainerParagraph0_1.paragraph_format.line_spacing = Inches(0.01)
farmInfoContainerParagraph0_1.paragraph_format.space_before = Pt(0)
farmInfoContainerParagraph0_1.paragraph_format.space_after = Pt(0)

# 농가현황(두수현황) table 생성
CowCountBoxWidth = [Cm(4), Cm(9.5)]
CowCountBoxHeight = [Cm(2), Cm(1), Cm(1), Cm(1)]
CowCountBox = makeTable(fInfocontainerCell0_1, row=4, col=2,
                        alignment='CENTER', width=CowCountBoxWidth, height=CowCountBoxHeight)

# 농가현황(두수현황) table 테두리 변경
titleBorder(
    CowCountBox,
    top_val='single',
    top_color='#F58D22',
    top_sz='5',
    bottom_val='single',
    bottom_color='#F58D22',
    bottom_sz='5',
    left_val='single',
    left_color='#F58D22',
    left_sz='5',
    right_val='single',
    right_color='#F58D22',
    right_sz='5'
)

# 농가현황(두수현황) table 배경색 변경
cellBackColor(CowCountBox, row=0, cell=0, color='FDECDA')
cellBackColor(CowCountBox, row=0, cell=1, color='FDECDA')
cellBackColor(CowCountBox, row=3, cell=0, color='FDECDA')
cellBackColor(CowCountBox, row=3, cell=1, color='FDECDA')

# 농가현황(두수현황) table index 입력
insertTextCell(CowCountBox, 0, 0, '구분', '000000', fontSize=12, fontStyle='bold',
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
insertTextCell(CowCountBox, 0, 1, '유전체두수', '000000', fontSize=12, fontStyle='bold',
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
insertTextCell(CowCountBox, 1, 0, '수', '000000', fontSize=12, fontStyle='bold',
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
insertTextCell(CowCountBox, 2, 0, '암', '000000', fontSize=12, fontStyle='bold',
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
insertTextCell(CowCountBox, 3, 0, '합계', '000000', fontSize=12, fontStyle='bold',
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

# 농가현황(두수현황) table data 입력
# valation => femaleCowCount, maleCowCount, unknownCowCount, allCount 참조

# 두수현황 - 수
insertTextCell(CowCountBox, 1, 1, str(maleCowCount), '000000', fontSize=12, fontStyle=None,
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
# 두수현황 - 암
insertTextCell(CowCountBox, 2, 1, str(femaleCowCount), '000000', fontSize=12, fontStyle=None,
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
# 두수현황 - unknown
if unknownCowCount == 0:
    insertTextCell(CowCountBox, 3, 1, str(allCount), '000000', fontSize=12, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
else:
    insertTextCell(CowCountBox, 3, 1, str(allCount), '000000', fontSize=12, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(CowCountBox, 3, 1, f'  (미확인 {str(unknownCowCount)}두 포함)', '000000', fontSize=7, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

lineSpace(doc, inches=0.1, space_before=0, space_after=0)

# 02.개월령 별 개체 분포도 headerBox 생성
makeHeaderBox(No="02", text="농가 개월령 별 개체 분포도")

lineSpace(doc, inches=0.1, space_before=0, space_after=0)

# 개월령 별 개체 분포도 이미지 table 생성
monthBoxWidth = [Cm(13.4), Cm(13.4)]
monthBoxHeight = [Cm(8), Cm(0.7)]
monthBox = makeTable(doc, row=2, col=2, alignment='CENTER',
                     width=monthBoxWidth, height=monthBoxHeight)

# 개월령 별 개체 분포도 이미지 table 테두리 변경
titleBorder(
    monthBox,
    top_val='single',
    top_color='#F58D22',
    top_sz='5',
    bottom_val='single',
    bottom_color='#F58D22',
    bottom_sz='5',
    left_val='single',
    left_color='#F58D22',
    left_sz='5',
    right_val='single',
    right_color='#F58D22',
    right_sz='5'
)

# 개월령 별 개체 분포도 이미지 table 배경색 변경
cellBackColor(monthBox, row=1, cell=0, color='FDECDA')
cellBackColor(monthBox, row=1, cell=1, color='FDECDA')
# 개월령 별 개체 분포도 이미지 table index 변경
insertTextCell(monthBox, 1, 0, '개월령 분포도 [수]', '000000', fontSize=12, fontStyle='bold',
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
insertTextCell(monthBox, 1, 1, '개월령 분포도 [암]', '000000', fontSize=12, fontStyle='bold',
               vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

# 개월령 별 개체 분포도 이미지 입력을 위한 cell 파라미터 지정
monthBoxparagraph0_0 = monthBox.rows[0].cells[0].paragraphs[0]
monthBoxparagraph0_1 = monthBox.rows[0].cells[1].paragraphs[0]

# 개월령 별 개체 분포도 이미지 삽입
add_float_picture(monthBoxparagraph0_0, './chart/monthchart_수.jpg',
                  width=Cm(12), height=Cm(7.5), pos_x=Cm(0.6), pos_y=Cm(0.3))
add_float_picture(monthBoxparagraph0_1, './chart/monthchart_암.jpg',
                  width=Cm(12), height=Cm(7.5), pos_x=Cm(0.6), pos_y=Cm(0.3))
lineSpace(doc, inches=0.01, space_before=0, space_after=0)

# 개체별 유전능력 제작
SNPChipName = 'Illumina Hanwoo 50K Beadchip'

# 전체 유전능력 data 생성
referenceSbvMean = referenceSbvMean.round(4)
referenceSbvMean = referenceSbvMean.to_list()

# 개체별 유전능력 data 생성
idvPageData = idvCowData.copy(deep=True)
idvPageData = idvPageData.astype({
    'BIR': 'string'
})
idvPageData = idvPageData.to_numpy()

for idv in idvPageData:
    # 03.개체별 유전능력 headerBox 생성
    makeHeaderBox(No="03", text="개체별 유전능력")

    lineSpace(doc, inches=0.1, space_before=0, space_after=0)

    idvCowInfoBoxWidth = [Cm(13), Cm(3.5), Cm(3.5), Cm(3.5), Cm(3.5)]
    idvCowInfoBoxHeight = [Cm(0.8), Cm(0.8), Cm(0.8), Cm(
        0.8), Cm(0.8), Cm(0.8), Cm(0.8), Cm(0.8), Cm(0.8)]
    idvCowInfoBox = makeTable(doc, row=9, col=5, alignment='CENTER',
                              width=idvCowInfoBoxWidth, height=idvCowInfoBoxHeight)

    titleBorder(
        idvCowInfoBox,
        top_val='single',
        top_color='#F58D22',
        top_sz='5',
        bottom_val='single',
        bottom_color='#F58D22',
        bottom_sz='5',
        left_val='single',
        left_color='#F58D22',
        left_sz='5',
        right_val='single',
        right_color='#F58D22',
        right_sz='5'
    )

    # 개체별 유전능력 data 입력
    insertTextCell(idvCowInfoBox, 8, 0, '유전능력', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    idvBoxParagraph0_0 = idvCowInfoBox.rows[0].cells[0].paragraphs[0]
    add_float_picture(idvBoxParagraph0_0, f'./chart/{idv[0]}_rader.jpg', width=Cm(
        6), height=Cm(6), pos_x=Cm(3.5), pos_y=Cm(0.2))

    insertTextCell(idvCowInfoBox, 0, 1, '개체번호', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 0, 2, idv[0], '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

    insertTextCell(idvCowInfoBox, 1, 1, 'SNP chip', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 1, 2, SNPChipName, '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

    insertTextCell(idvCowInfoBox, 2, 1, '생년월일', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 2, 2, idv[1], '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

    insertTextCell(idvCowInfoBox, 2, 3, '성별', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 2, 4, idv[2], '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

    insertTextCell(idvCowInfoBox, 3, 1, '친자확인결과', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 3, 2, idv[3], '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

    insertTextCell(idvCowInfoBox, 3, 3, '친자정확도', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 3, 4, idv[4], '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

    insertTextCell(idvCowInfoBox, 4, 1, '형질', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 5, 1, '도체중-CWT', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 6, 1, '등심단면적-EMA', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 7, 1, '등지방두께-BFT', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 8, 1, '근내지방-MAR', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

    insertTextCell(idvCowInfoBox, 4, 2, '전체 평균', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 5, 2, str(referenceSbvMean[0]), '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 6, 2, str(referenceSbvMean[1]), '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 7, 2, str(referenceSbvMean[2]), '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 8, 2, str(referenceSbvMean[3]), '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

    insertTextCell(idvCowInfoBox, 4, 3, '개체 능력', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 5, 3, str(idv[5]), '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 6, 3, str(idv[6]), '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 7, 3, str(idv[7]), '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 8, 3, str(idv[8]), '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

    insertTextCell(idvCowInfoBox, 4, 4, '등급', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 5, 4, str(idv[9]), '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 6, 4, str(idv[10]), '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 7, 4, str(idv[11]), '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    insertTextCell(idvCowInfoBox, 8, 4, str(idv[12]), '000000', fontSize=9, fontStyle=None,
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))

    # idvCowInfoBox 배경색 변경
    cellBackColor(idvCowInfoBox, row=8, cell=0, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=0, cell=1, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=1, cell=1, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=2, cell=1, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=3, cell=1, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=4, cell=1, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=5, cell=1, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=6, cell=1, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=7, cell=1, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=8, cell=1, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=2, cell=3, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=3, cell=3, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=4, cell=2, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=4, cell=3, color='FDECDA')
    cellBackColor(idvCowInfoBox, row=4, cell=4, color='FDECDA')

    # idvCowInfoBox table cell 병합
    idvCellList1 = [0, 0]
    cellList1 = [[1, 0], [2, 0], [3, 0], [4, 0], [5, 0], [6, 0], [7, 0]]
    cellMerge(idvCowInfoBox, idvCellList1, cellList1)

    idvCellList2 = [0, 2]
    cellList2 = [[0, 3], [0, 4]]
    cellMerge(idvCowInfoBox, idvCellList2, cellList2)

    idvCellList3 = [1, 2]
    cellList3 = [[1, 3], [1, 4]]
    cellMerge(idvCowInfoBox, idvCellList3, cellList3)

    # paraTitle 제작
    lineSpace(doc, inches=0.1, space_before=0, space_after=0)

    paraTitle = doc.add_paragraph()
    paraTitle.paragraph_format.left_indent = Inches(0.1)
    paraTitle.paragraph_format.space_before = Pt(0)
    paraTitle.paragraph_format.space_after = Pt(0)
    paragraphText(paraTitle, text='▶', fontsize=16,
                  color='F58D22', alignment=NONE, style='bold')
    paragraphText(paraTitle, text=' 형질별 개체 순위 백분율', fontsize=16,
                  color='000000', alignment=NONE, style='bold')

    # perRankBox 생성
    perRankBoxWidth = [Cm(6.75), Cm(6.75), Cm(6.75), Cm(6.75)]
    perRankBoxHeight = [Cm(0.5), Cm(7.5)]
    perRankBox = makeTable(doc, row=2, col=4, alignment='CENTER',
                           width=perRankBoxWidth, height=perRankBoxHeight)

    titleBorder(
        perRankBox,
        top_val='single',
        top_color='#F58D22',
        top_sz='5',
        bottom_val='single',
        bottom_color='#F58D22',
        bottom_sz='5',
        left_val='single',
        left_color='#F58D22',
        left_sz='5',
        right_val='single',
        right_color='#F58D22',
        right_sz='5'
    )

    # perRankBox data, image 삽입
    insertTextCell(perRankBox, 0, 0, '도체중', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    perRankBoxParagraph1_1 = perRankBox.rows[1].cells[0].paragraphs[0]
    add_float_picture(perRankBoxParagraph1_1, f'./chart/{idv[0]}_CWT_pchart.jpg', width=Cm(
        3.5), height=Cm(7), pos_x=Cm(1.5), pos_y=Cm(0.3))

    insertTextCell(perRankBox, 0, 1, '등심단면적', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    perRankBoxParagraph1_2 = perRankBox.rows[1].cells[1].paragraphs[0]
    add_float_picture(perRankBoxParagraph1_2, f'./chart/{idv[0]}_EMA_pchart.jpg', width=Cm(
        3.5), height=Cm(7), pos_x=Cm(1.5), pos_y=Cm(0.3))

    insertTextCell(perRankBox, 0, 2, '등지방두께', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    perRankBoxParagraph1_3 = perRankBox.rows[1].cells[2].paragraphs[0]
    add_float_picture(perRankBoxParagraph1_3, f'./chart/{idv[0]}_BFT_pchart.jpg', width=Cm(
        3.5), height=Cm(7), pos_x=Cm(1.5), pos_y=Cm(0.3))

    insertTextCell(perRankBox, 0, 3, '근내지방', '000000', fontSize=9, fontStyle='bold',
                   vertical_alignment='CENTER', para_alignment='CENTER', space_after=Pt(0), space_before=Pt(0))
    perRankBoxParagraph1_4 = perRankBox.rows[1].cells[3].paragraphs[0]
    add_float_picture(perRankBoxParagraph1_4, f'./chart/{idv[0]}_MAR_pchart.jpg', width=Cm(
        3.5), height=Cm(7), pos_x=Cm(1.5), pos_y=Cm(0.3))

    cellBackColor(perRankBox, row=0, cell=0, color='FDECDA')
    cellBackColor(perRankBox, row=0, cell=1, color='FDECDA')
    cellBackColor(perRankBox, row=0, cell=2, color='FDECDA')
    cellBackColor(perRankBox, row=0, cell=3, color='FDECDA')
    lineSpace(doc, inches=0.01, space_before=0, space_after=0)


# 04.선발지수 적용 농가 개체 리스트 headerBox 생성
makeHeaderBox(No="04", text="선발지수 적용 농가 개체 리스트")
lineSpace(doc, inches=0.01, space_before=0, space_after=0)
paraTitle = doc.add_paragraph()
paraTitle.paragraph_format.left_indent = Inches(0.1)
paraTitle.paragraph_format.space_before = Pt(0)
paraTitle.paragraph_format.space_after = Pt(0)
paragraphText(paraTitle, text='▶', fontsize=13,
              color='F58D22', alignment=NONE, style='bold')
paragraphText(paraTitle, text=' 소득최적선발식(가중치): (도체중*7) + (등심단면적*3) + (등지방두께*-3) + (근내지방도*4)',
              fontsize=13, color='000000', alignment=NONE, style='bold')
lineSpace(doc, inches=0.01, space_before=0, space_after=0)

# indexRankBox 생성
indexRankBoxWidth = [Cm(3.8), Cm(3.8), Cm(3.8), Cm(1.6),
                     Cm(3), Cm(3), Cm(3), Cm(3), Cm(2)]
indexRankBox = makeTable(
    doc, row=selectIndexData.shape[0]+1, col=selectIndexData.shape[1], alignment='CENTER', width=indexRankBoxWidth, height=None)

# indexRankBox 헤더 추가
for j in range(selectIndexData.shape[-1]):
    insertTextCell(indexRankBox, 0, j, selectIndexData.columns[j], color='000000', fontSize=11, fontStyle='bold',
                   vertical_alignment="CENTER", para_alignment="CENTER", space_after=Pt(1), space_before=Pt(1))
    indexRankBox.rows[0].height = Cm(2)

# indexRankBox 데이터 추가
for i in range(selectIndexData.shape[0]):
    for j in range(selectIndexData.shape[-1]):
        insertTextCell(indexRankBox, i+1, j, str(selectIndexData.values[i, j]), color='000000', fontSize=10,
                       fontStyle=None, vertical_alignment="CENTER", para_alignment="CENTER", space_after=Pt(1), space_before=Pt(1))
        indexRankBox.rows[i+1].height = Cm(1)

# indexRankBox header 반복처리
set_repeat_table_header(indexRankBox.rows[0])

titleBorder(
    indexRankBox,
    top_val='single',
    top_color='#F58D22',
    top_sz='5',
    bottom_val='single',
    bottom_color='#F58D22',
    bottom_sz='5',
    left_val='single',
    left_color='#F58D22',
    left_sz='5',
    right_val='single',
    right_color='#F58D22',
    right_sz='5'
)

# indexRankBox 배경색 설정
cellBackColor(indexRankBox, row=0, cell=0, color='FDECDA')
cellBackColor(indexRankBox, row=0, cell=1, color='FDECDA')
cellBackColor(indexRankBox, row=0, cell=2, color='FDECDA')
cellBackColor(indexRankBox, row=0, cell=3, color='FDECDA')
cellBackColor(indexRankBox, row=0, cell=4, color='FDECDA')
cellBackColor(indexRankBox, row=0, cell=5, color='FDECDA')
cellBackColor(indexRankBox, row=0, cell=6, color='FDECDA')
cellBackColor(indexRankBox, row=0, cell=7, color='FDECDA')
cellBackColor(indexRankBox, row=0, cell=8, color='FDECDA')

########################################## 보고서 완료##########################################보고서 완료##########################################보고서 완료##########################################보고서 완료##########################################

inputFile = f'./word/{farmNameReport}_유전체보고서.docx'

doc.save(inputFile)

outputFile = f'./pdf/{farmNameReport}_유전체보고서.pdf'

file = open(outputFile, "w")
file.close()

print(f'./pdf/{farmNameReport}_유전체보고서.pdf 생성 중 .....')

convert(inputFile, outputFile)
